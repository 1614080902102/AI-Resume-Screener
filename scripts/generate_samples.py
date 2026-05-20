"""Bulk generate 75 unique candidates across 5 jobs.

Structure:
  examples/candidates/<job_slug>/<candidate_id>/{resume.pdf, cover.docx}
  examples/candidates/<job_slug>/<candidate_id>.zip
"""
from __future__ import annotations

import random
import shutil
import zipfile
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

ROOT = Path(__file__).parent.parent
OUT_DIR = ROOT / "examples" / "candidates"

random.seed(42)

# ============================================================
# Name + location pools
# ============================================================
FIRST_NAMES_M = [
    "博文", "宇翔", "明哲", "子涵", "浩然", "嘉豪", "皓轩", "昊然", "俊辰", "梓睿",
    "晨曦", "天佑", "煜城", "瑞峰", "辰逸", "鸿涛", "懿轩", "锦程", "卓尔", "尚阳",
]
FIRST_NAMES_F = [
    "雨菲", "思琪", "晓晨", "雪薇", "怡然", "婉清", "心怡", "梓萱", "雨欣", "诗涵",
    "佳怡", "若曦", "梦琪", "雅静", "依诺", "玥彤", "汐月", "佩瑶", "若宁", "知微",
]
LAST_NAMES = [
    "张", "李", "王", "刘", "陈", "杨", "黄", "周", "吴", "徐",
    "孙", "马", "朱", "胡", "郭", "何", "高", "林", "罗", "郑",
    "梁", "谢", "宋", "唐", "许", "邓", "韩", "冯", "曹", "彭",
]
LOCATIONS = ["北京", "上海", "广州", "深圳", "杭州", "成都", "南京", "苏州", "厦门", "西安"]


def gen_name(gender: str) -> str:
    last = random.choice(LAST_NAMES)
    first = random.choice(FIRST_NAMES_M if gender == "男" else FIRST_NAMES_F)
    return last + first


def gen_id(name: str, idx: int) -> str:
    return f"{idx:02d}_{name}"


# ============================================================
# PDF / Word helpers
# ============================================================
def make_pdf(path: Path, title: str, sections: list[tuple[str, str]]):
    doc = SimpleDocTemplate(str(path), pagesize=A4, topMargin=2 * cm)
    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]; h2 = styles["Heading2"]; body = styles["BodyText"]
    h1.fontName = h2.fontName = body.fontName = "STSong-Light"
    h1.fontSize = 16; h2.fontSize = 13; body.fontSize = 10; body.leading = 16
    story = [Paragraph(title, h1), Spacer(1, 0.3 * cm)]
    for heading, content in sections:
        story.append(Paragraph(heading, h2))
        story.append(Spacer(1, 0.15 * cm))
        for line in content.strip().split("\n"):
            story.append(Paragraph(line.replace("&", "&amp;"), body))
        story.append(Spacer(1, 0.25 * cm))
    doc.build(story)


def make_docx(path: Path, title: str, sections: list[tuple[str, str]]):
    doc = Document()
    doc.add_heading(title, level=0)
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        for line in content.strip().split("\n"):
            doc.add_paragraph(line)
    doc.save(str(path))


def zip_folder(src_dir: Path, zip_path: Path):
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in src_dir.iterdir():
            if f.is_file():
                zf.write(f, arcname=f.name)


# ============================================================
# 5 个岗位 × 15 个候选人 (5 强 / 5 中 / 5 弱)
# ============================================================
JOBS = {
    "ai_agent": {
        "name": "AI Agent 开发工程师",
        "strong": {
            "title": "AI 全栈应用工程师",
            "exp_template": [
                "{years} 年以上 AI 应用开发经验,主导 ToB 知识库 / Agent 项目",
                "FastAPI + RAG + Embedding 调优实战,中文召回率从 60%+ 优化至 85%+",
                "Dify / LangChain / MCP 协议熟练,多模型路由(GPT/Claude/DeepSeek/Ollama)",
                "复杂爬虫经验,Scrapy + Playwright + 反爬绕过",
            ],
        },
        "medium": {
            "title": "Python 后端工程师",
            "exp_template": [
                "{years} 年 Python 后端经验,FastAPI / Django 熟悉",
                "PostgreSQL 高并发场景调优,慢查询优化",
                "对 AI / LLM 有兴趣,业余跑过 LangChain demo",
                "没做过完整 Agent 产品,但工程基础扎实",
            ],
        },
        "weak": {
            "title": "前端工程师",
            "exp_template": [
                "{years} 年前端开发,Vue / React 实战",
                "学过一点 Python,做过爬虫小工具",
                "对 AI 感兴趣,玩过 ChatGPT API",
                "无 RAG / Agent 实战经验",
            ],
        },
    },
    "event_planning": {
        "name": "活动策划",
        "strong": {
            "title": "B 端市场活动负责人",
            "exp_template": [
                "{years} 年 B2B 活动策划,主导招商会 / 行业沙龙 / 海外推介",
                "操盘新加坡 / 曼谷 / 迪拜目的地推介会,单场客户转化 50+",
                "熟悉跨境 / 旅游 / 酒店行业,英语流利",
                "活动 ROI 可量化,有完整数据复盘体系",
            ],
        },
        "medium": {
            "title": "市场专员",
            "exp_template": [
                "{years} 年活动经验,主要做公司内部活动 + 校园招聘",
                "做过 1-2 次外部行业沙龙,数据一般",
                "对跨境业务感兴趣,但缺实战",
                "英语日常沟通可,正在备考雅思",
            ],
        },
        "weak": {
            "title": "市场实习生 / 校园活动负责人",
            "exp_template": [
                "实习 / 学生工作经验,做过校园招聘会",
                "组织过社团活动 / 公司年会",
                "无 B 端客户活动经验",
                "无国际业务背景",
            ],
        },
    },
    "overseas_marketing": {
        "name": "海外营销专员",
        "strong": {
            "title": "海外 B2B 增长营销",
            "exp_template": [
                "{years} 年跨境 B2B 营销,LinkedIn / WhatsApp / EDM 多渠道操盘",
                "英语流利(海外留学),熟悉欧美 / 东南亚市场",
                "用过 HubSpot / Salesforce / Brevo 等 CRM,自动化经验丰富",
                "跨境电商或酒店行业背景,文化敏感度高",
            ],
        },
        "medium": {
            "title": "数字营销 / 内容运营",
            "exp_template": [
                "{years} 年国内数字营销,SEM / 信息流投放熟练",
                "英语 CET6,基本沟通无障碍",
                "做过海外社媒账号但规模小",
                "未深度接触跨境 B2B 客户",
            ],
        },
        "weak": {
            "title": "跨境电商运营",
            "exp_template": [
                "{years} 年跨境电商运营(亚马逊 / 速卖通)",
                "主要做 ToC 选品 + 上架,不熟 B2B",
                "英语阅读可,口语一般",
                "无 LinkedIn / EDM 营销经验",
            ],
        },
    },
    "frontend": {
        "name": "前端工程师 (Vue)",
        "strong": {
            "title": "高级前端工程师",
            "exp_template": [
                "{years} 年 Vue 实战(Vue 2 → Vue 3 迁移经验)",
                "TypeScript / Vite / Pinia / Vue Router 精通",
                "B 端中后台大型项目,ECharts / AntV 数据可视化",
                "接触过 AI 应用前端(对话 UI / 流式输出)",
            ],
        },
        "medium": {
            "title": "前端工程师",
            "exp_template": [
                "{years} 年 React 经验,正在转 Vue",
                "TypeScript 熟悉,Pinia 学过 demo",
                "C 端项目为主,中后台经验少",
                "对 AI 前端有了解但未实战",
            ],
        },
        "weak": {
            "title": "前端实习 / 初级",
            "exp_template": [
                "实习 / 1-2 年前端经验,主要写页面",
                "Vue 用过基础语法,无项目主导经验",
                "TypeScript 学过",
                "无 B 端中后台经验",
            ],
        },
    },
    "ai_content": {
        "name": "AI 内容运营",
        "strong": {
            "title": "AI 内容运营 / Prompt 工程师",
            "exp_template": [
                "{years} 年内容运营,深度使用 ChatGPT / Claude 做营销文案",
                "公众号 / 小红书 / Newsletter 操盘,粉丝 10w+",
                "Prompt 工程实战,熟悉 AI 文案的迭代优化方法",
                "数据敏感,基于 CTR / 转化率持续优化",
            ],
        },
        "medium": {
            "title": "内容运营",
            "exp_template": [
                "{years} 年内容运营,公众号 / 短视频脚本",
                "用过 ChatGPT 但停留在 demo 层",
                "数据看得懂但优化经验少",
                "对 Prompt 工程感兴趣",
            ],
        },
        "weak": {
            "title": "新媒体编辑",
            "exp_template": [
                "{years} 年新媒体编辑,主要排版 + 转载",
                "ChatGPT 偶尔用,无系统 Prompt 经验",
                "数据更多是堆数字汇报",
                "AI 工具掌握度低",
            ],
        },
    },
}

EDUCATION_STRONG = [
    "2014 - 2018 | 浙江大学 | 计算机科学与技术 | 本科",
    "2015 - 2019 | 复旦大学 | 软件工程 | 本科",
    "2016 - 2020 | 北京邮电大学 | 计算机科学 | 本科",
    "2014 - 2018 | 中山大学 | 计算机科学 | 本科",
    "2015 - 2019 | 华中科技大学 | 软件工程 | 本科",
]
EDUCATION_MID = [
    "2016 - 2020 | 北京工业大学 | 软件工程 | 本科",
    "2015 - 2019 | 西安电子科技大学 | 通信工程 | 本科",
    "2017 - 2021 | 暨南大学 | 计算机 | 本科",
    "2016 - 2020 | 南京邮电大学 | 软件 | 本科",
    "2015 - 2019 | 江苏大学 | 计算机科学 | 本科",
]
EDUCATION_WEAK = [
    "2018 - 2022 | 某二本院校 | 计算机 | 本科",
    "2019 - 2023 | 某三本院校 | 软件工程 | 本科",
    "2018 - 2021 | 某专科 | 计算机应用 | 大专",
    "2019 - 2023 | 某独立学院 | 软件 | 本科",
    "2017 - 2021 | 某地方院校 | 计算机 | 本科",
]


def gen_one(job_slug: str, tier: str, idx_in_tier: int, global_idx: int) -> dict:
    job = JOBS[job_slug]
    job_name = job["name"]
    profile = job[tier]
    gender = random.choice(["男", "女"])
    name = gen_name(gender)
    # 经验年限按 tier
    years_map = {"strong": [4, 5, 6, 7, 8], "medium": [2, 3, 4, 5], "weak": [0, 1, 2]}
    years = random.choice(years_map[tier])
    age = 22 + years + random.randint(1, 3)
    location = random.choice(LOCATIONS)

    edu_pool = {"strong": EDUCATION_STRONG, "medium": EDUCATION_MID, "weak": EDUCATION_WEAK}
    edu = random.choice(edu_pool[tier])

    # 经历摘要
    exp_lines = "\n".join(
        line.format(years=years) for line in profile["exp_template"]
    )

    # cover letter intro
    why_lines = {
        "strong": f"我对贵司{job_name}岗位特别感兴趣 — 我过去 {years} 年的经历高度契合这条线,希望能直接上手。",
        "medium": f"我希望从现有方向转向{job_name}领域 — 基础扎实,愿意快速补齐缺口。",
        "weak": f"我对{job_name}方向充满热情,虽然经验有限,但学习能力强,希望加入贵司从基础做起。",
    }[tier]

    return {
        "name": name,
        "gender": gender,
        "age": age,
        "location": location,
        "years": years,
        "edu": edu,
        "title": profile["title"],
        "exp": exp_lines,
        "why": why_lines,
        "tier": tier,
        "phone": f"1{random.choice('3578')}{random.randint(0, 9)}-{random.randint(1000, 9999)}-{random.randint(1000, 9999)}",
        "email": f"{name.lower()[:6]}.{random.randint(100, 999)}@example.com",
    }


def write_candidate(out_root: Path, job_slug: str, cand: dict, idx: int):
    cand_id = gen_id(cand["name"], idx)
    cand_dir = out_root / job_slug / cand_id
    cand_dir.mkdir(parents=True, exist_ok=True)

    # PDF resume
    make_pdf(
        cand_dir / f"{cand_id}_简历.pdf",
        f"{cand['name']} - {cand['title']}",
        [
            (
                "个人信息",
                f"性别:{cand['gender']} | 年龄:{cand['age']} | 现居:{cand['location']}\n"
                f"电话:{cand['phone']} | 邮箱:{cand['email']}\n"
                f"工作年限:{cand['years']} 年",
            ),
            ("教育经历", cand["edu"]),
            ("工作经历与项目", cand["exp"]),
        ],
    )

    # Word cover letter
    make_docx(
        cand_dir / f"{cand_id}_自荐信.docx",
        f"求职自荐信 - {cand['name']}",
        [
            ("求职动机", cand["why"]),
        ],
    )

    # Zip
    zip_folder(cand_dir, out_root / job_slug / f"{cand_id}.zip")


def main():
    if OUT_DIR.exists():
        # 清理旧数据
        for old in OUT_DIR.iterdir():
            if old.is_dir() or old.suffix == ".zip":
                if old.is_dir():
                    shutil.rmtree(old)
                else:
                    old.unlink()
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    total = 0
    global_idx = 0
    for job_slug in JOBS:
        for tier_idx, tier in enumerate(["strong", "medium", "weak"]):
            for k in range(5):
                global_idx += 1
                cand = gen_one(job_slug, tier, k, global_idx)
                write_candidate(OUT_DIR, job_slug, cand, global_idx)
                total += 1
        print(f"  ✓ {JOBS[job_slug]['name']:24s} - 15 candidates")
    print(f"\n总共生成 {total} 个候选人,目录:{OUT_DIR}")


if __name__ == "__main__":
    main()
